from __future__ import annotations
from typing import List, Dict, Any
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pathlib

# -------------------------------------------------------------
# GLOBAL accumulator (used only if pipeline does not pass final_exam)
# -------------------------------------------------------------
FINAL_EXAM_QUESTIONS: List[Dict[str, Any]] = []


# -------------------------------------------------------------
# Styles
# -------------------------------------------------------------
def ensure_styles(doc: Document) -> None:
    styles = doc.styles
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        if style_name not in styles:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)


# -------------------------------------------------------------
# Helpers for clean rendering
# -------------------------------------------------------------
def _add_image_placeholder(doc: Document, image_id: str | None):
    """Render a simple placeholder for image_id."""
    if image_id:
        p = doc.add_paragraph()
        r = p.add_run(f"[Image: {image_id}]")
        r.italic = True


def _add_button_label(doc: Document, label: str | None):
    """Show button label visually."""
    if label:
        p = doc.add_paragraph()
        run = p.add_run(f"[Button Label: {label}]")
        run.bold = True


def _add_spacing(doc: Document, lines: int = 1):
    for _ in range(lines):
        doc.add_paragraph("")


# -------------------------------------------------------------
# Add ONE block (header + pages + inline quiz)
# -------------------------------------------------------------
def add_block_to_doc(doc: Document, block: Dict[str, Any]) -> None:
    header = block.get("header", "Untitled Section")
    pages: List[Dict[str, Any]] = block.get("pages", [])
    quiz: List[Dict[str, Any]] = block.get("quiz", [])

    # ---------------------------------------------------------
    # Header
    # ---------------------------------------------------------
    h1 = doc.add_heading(header, level=1)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # ---------------------------------------------------------
    # Pages
    # ---------------------------------------------------------
    for page in pages:
        p_type = page.get("type", "page")

        # --------------------------------------------------
        # PAGE BLOCK
        # --------------------------------------------------
        if p_type == "page":
            title = page.get("title", "")
            if title:
                doc.add_heading(title, level=2)

            _add_image_placeholder(doc, page.get("image_id"))
            doc.add_paragraph(page.get("content", ""))
            _add_spacing(doc)

        # --------------------------------------------------
        # ENGAGE BLOCK
        # --------------------------------------------------
        elif p_type == "engage":
            intro = page.get("intro", {})
            items = page.get("items", [])

            # Intro
            doc.add_heading("(Engage Intro)", level=2)
            _add_image_placeholder(doc, intro.get("image_id"))
            doc.add_paragraph(intro.get("content", ""))
            _add_spacing(doc)

            # Items
            for idx, item in enumerate(items, start=1):
                doc.add_heading(f"Engage Item {idx}", level=3)
                _add_button_label(doc, item.get("button_label"))
                _add_image_placeholder(doc, item.get("image_id"))
                doc.add_paragraph(item.get("content", ""))
                _add_spacing(doc)

        # --------------------------------------------------
        # ENGAGE2 BLOCK
        # --------------------------------------------------
        elif p_type == "engage2":
            intro = page.get("intro", {})
            steps = page.get("steps", [])

            # Intro
            doc.add_heading("(Engage 2 Intro)", level=2)
            _add_image_placeholder(doc, intro.get("image_id"))
            doc.add_paragraph(intro.get("content", ""))
            _add_spacing(doc)

            # Steps (each step may have its own label)
            for idx, step in enumerate(steps, start=1):
                doc.add_heading(f"Engage 2 Step {idx}", level=3)
                _add_button_label(doc, step.get("button_label"))
                _add_image_placeholder(doc, step.get("image_id"))
                doc.add_paragraph(step.get("content", ""))
                _add_spacing(doc)

        else:
            # Unknown type fallback
            doc.add_paragraph(f"[Unknown block type: {p_type}]")
            _add_spacing(doc)

    # ---------------------------------------------------------
    # Inline Quiz
    # ---------------------------------------------------------
    inline_questions = []
    reserved_questions = []

    for q in quiz:
        if q.get("reserve_for_final_exam"):
            reserved_questions.append(q)
        else:
            inline_questions.append(q)

    # Store reserved for later
    for rq in reserved_questions:
        FINAL_EXAM_QUESTIONS.append({
            "header": header,
            **rq
        })

    # Render inline quiz
    if inline_questions:
        doc.add_heading(f"Quiz: {header}", level=3)

        for idx, q in enumerate(inline_questions, start=1):
            question = q.get("question", "")
            options = q.get("options", [])
            correct = q.get("correct_answers", [])
            qtype = q.get("type", "")

            doc.add_paragraph(f"Q{idx}. {question} ({qtype})")

            label_letters = "abcdefghijklmnopqrstuvwxyz"
            for opt_idx, opt in enumerate(options):
                doc.add_paragraph(f"   {label_letters[opt_idx]}) {opt}")

            if correct:
                p = doc.add_paragraph()
                p.add_run(f"Correct: {', '.join(correct)}").bold = True

            _add_spacing(doc, 1)


# -------------------------------------------------------------
# Final Exam Section
# -------------------------------------------------------------
def add_final_exam_section(doc: Document, final_exam_list: List[Dict[str, Any]]) -> None:
    if not final_exam_list:
        return

    doc.add_page_break()
    doc.add_heading("Final Exam", level=1)

    for idx, q in enumerate(final_exam_list, start=1):
        question = q.get("question", "")
        options = q.get("options", [])
        correct = q.get("correct_answers", [])
        qtype = q.get("type", "")
        source = q.get("header", "")

        doc.add_paragraph(f"Q{idx}. {question} ({qtype}) — from {source}")

        label_letters = "abcdefghijklmnopqrstuvwxyz"
        for opt_idx, opt in enumerate(options):
            doc.add_paragraph(f"   {label_letters[opt_idx]}) {opt}")

        if correct:
            p = doc.add_paragraph()
            p.add_run(f"Correct: {', '.join(correct)}").bold = True

        _add_spacing(doc)


# -------------------------------------------------------------
# Build final DOCX
# -------------------------------------------------------------
def build_docx_from_stage2(stage2_dir, output_path, module_title=None, final_exam=None) -> None:
    import json
    import pathlib

    FINAL_EXAM_QUESTIONS.clear()

    doc = Document()
    ensure_styles(doc)

    # Title page
    if module_title:
        p = doc.add_paragraph()
        r = p.add_run(module_title)
        r.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()

    # Process blocks
    for path in sorted(pathlib.Path(stage2_dir).glob("*.json")):
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)

        add_block_to_doc(doc, data)
        doc.add_page_break()

    final_exam_list = final_exam if final_exam is not None else FINAL_EXAM_QUESTIONS

    add_final_exam_section(doc, final_exam_list)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"DOCX written to: {output_path}")
