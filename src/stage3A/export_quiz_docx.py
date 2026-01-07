# src/stage3A/export_quiz_docx.py
from __future__ import annotations

import json
from pathlib import Path
from typing import List, Dict, Any
from collections import defaultdict

from docx import Document
from docx.shared import Pt


INPUT_PATH = Path("data/processed/quiz_flattened_stage2_9_2.json")
OUTPUT_PATH = Path("data/exports/quiz_questions.docx")

PLACEMENT_ORDER = ["inline", "application", "final"]


def load_questions(path: Path) -> List[Dict[str, Any]]:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def create_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


def write_quiz_header(doc: Document, quiz_id: int) -> None:
    doc.add_heading(f"Quiz {quiz_id}", level=1)


def write_question(
    doc: Document,
    index: int,
    question: Dict[str, Any],
) -> None:
    doc.add_paragraph(f"Question {index}", style="Heading 3")
    doc.add_paragraph(question["prompt"])

    q_type = question.get("type") or question.get("question_type")

    if q_type == "mcq":
        for key in ["A", "B", "C", "D"]:
            opt_key = f"option_{key}"
            if opt_key not in question:
                raise KeyError(
                    f"Missing {opt_key} in question {question.get('question_id')}"
                )
            doc.add_paragraph(
                f"{key}. {question[opt_key]}",
                style="List Bullet",
            )

        doc.add_paragraph(f"Correct Answer: {question['correct_answer']}")

    elif q_type == "true_false":
        answer = "True" if question["correct_answer"] is True else "False"
        doc.add_paragraph(f"Correct Answer: {answer}")

    else:
        raise ValueError(f"Unknown question type: {q_type}")

    doc.add_paragraph("Rationale:")
    doc.add_paragraph(question["rationale"])
    doc.add_paragraph("")


def export_quizzes_to_docx(
    questions: List[Dict[str, Any]],
    output_path: Path,
) -> None:
    doc = create_document()

    # Group questions by quiz_id
    quizzes: Dict[int, List[Dict[str, Any]]] = defaultdict(list)
    for q in questions:
        quizzes[q["quiz_id"]].append(q)

    for quiz_id in sorted(quizzes.keys()):
        quiz_questions = quizzes[quiz_id]

        write_quiz_header(doc, quiz_id)

        # Group by placement
        by_placement: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
        for q in quiz_questions:
            by_placement[q.get("placement", "unknown")].append(q)

        question_counter = 1

        for placement in PLACEMENT_ORDER:
            section_questions = by_placement.get(placement)
            if not section_questions:
                continue

            # Placement-specific section titles
            if placement == "inline":
                section_title = f"Inline Quiz {quiz_id}"
            elif placement == "application":
                section_title = "Application Questions"
            elif placement == "final":
                section_title = "Final Questions"
            else:
                section_title = placement.title()

            doc.add_heading(section_title, level=2)

            for q in section_questions:
                write_question(doc, question_counter, q)
                question_counter += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    questions = load_questions(INPUT_PATH)
    export_quizzes_to_docx(questions, OUTPUT_PATH)
    print(f"[Stage 3A] Quiz questions exported to {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
