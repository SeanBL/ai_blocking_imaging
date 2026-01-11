from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
import sys


RAW_DIR = Path("data/raw")  # ← adjust if needed


def inspect_docx(path: Path) -> None:
    doc = Document(path)

    print(f"\nInspecting: {path}\n")

    for ti, table in enumerate(doc.tables):
        print(f"\n=== TABLE {ti} ===")

        for ri, row in enumerate(table.rows):
            print(f"\n-- Row {ri} --")

            for ci, cell in enumerate(row.cells):
                print(f"\nCell {ci}:")

                for pi, p in enumerate(cell.paragraphs):
                    text = p.text.strip()
                    style = p.style.name if p.style else "None"

                    # Check numbering (bullets / lists)
                    numPr = p._p.find(qn("w:numPr"))
                    has_numbering = numPr is not None

                    print(
                        f"  P{pi}: "
                        f"text='{text}' | "
                        f"style='{style}' | "
                        f"numbered={has_numbering}"
                    )


def resolve_input_path(argv) -> Path:
    # Case 1: explicit path provided
    if len(argv) == 2:
        path = Path(argv[1])
        if not path.exists():
            raise SystemExit(f"❌ File not found: {path}")
        if path.suffix.lower() != ".docx":
            raise SystemExit("❌ Input file must be a .docx")
        return path

    # Case 2: auto-scan directory
    if len(argv) == 1:
        docs = list(RAW_DIR.glob("*.docx"))

        if not docs:
            raise SystemExit(f"❌ No .docx files found in {RAW_DIR}")
        if len(docs) > 1:
            raise SystemExit(
                f"❌ Multiple .docx files found in {RAW_DIR}:\n"
                + "\n".join(f"  - {d.name}" for d in docs)
            )

        return docs[0]

    raise SystemExit(
        "Usage:\n"
        "  python -m src.debug.inspect_word_structure [optional_docx_path]"
    )


def main(argv):
    path = resolve_input_path(argv)
    inspect_docx(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
