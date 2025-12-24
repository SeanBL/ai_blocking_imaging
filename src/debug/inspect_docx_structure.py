from docx import Document
from pathlib import Path

doc = Document("data/raw/1561_Advanced_Training_Child_Health.docx")

print(f"PARAGRAPHS: {len(doc.paragraphs)}")
for i, p in enumerate(doc.paragraphs[:20], 1):
    print(f"P{i:02d}: {repr(p.text)}")

print("\nTABLES:", len(doc.tables))
for ti, tbl in enumerate(doc.tables, 1):
    print(f"\n=== TABLE {ti} ===")
    for ri, row in enumerate(tbl.rows, 1):
        cells = [c.text.strip().replace('\n', ' | ') for c in row.cells]
        print(f"Row {ri}: {cells}")
