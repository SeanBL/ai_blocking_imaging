from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Iterable, List

from docx import Document

from .normalization import normalize_and_filter

def _is_column_header_row(cells: list[str]) -> bool:
    """
    Detects Word table column header rows that should NOT be audited.

    This is intentionally strict and explicit.
    """
    normalized = [c.lower() for c in cells]

    KNOWN_HEADERS = {
        "image",
        "english text",
        "notes and instructions",
        "button labels",
    }

    hits = sum(1 for c in normalized if c in KNOWN_HEADERS)
    return hits >= 2

# -------------------------------------------------
# WORD DOCUMENT TEXT EXTRACTION
# -------------------------------------------------

def extract_word_text_fragments(docx_path: Path) -> List[str]:
    """
    Extract all meaningful text fragments from Word slide tables ONLY.

    Stage 2.1 audits text that Stage 1 is structurally responsible for.
    Free-floating paragraphs outside slide tables are intentionally ignored.
    """

    doc = Document(str(docx_path))
    raw_fragments: List[str] = []

    for tbl in doc.tables:
        for row in tbl.rows:
            cell_texts = []

            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = p.text.strip()
                    if txt:
                        cell_texts.append(txt)

            # Skip column header rows (structural labels, not content)
            if _is_column_header_row(cell_texts):
                continue

            # Everything else inside slide tables IS valid content
            raw_fragments.extend(cell_texts)

    # Split Stage 2 fragments by lines to match Stage 1 granularity
    split_fragments: List[str] = []

    for frag in raw_fragments:
        if not isinstance(frag, str):
            continue
        for part in frag.splitlines():
            part = part.replace("\u00A0", " ")
            part = " ".join(part.split()).strip()
            if part:
                split_fragments.append(part)

    return normalize_and_filter(split_fragments)



# -------------------------------------------------
# STAGE 2 JSON TEXT EXTRACTION
# -------------------------------------------------

def extract_stage2_text_fragments(module: Dict[str, Any]) -> List[str]:
    """
    Extract all meaningful text fragments from Stage 2 output.

    Traverses:
    - module_title
    - slide headers
    - notes
    - images (textual identifiers)
    - panel content blocks
    - engage1 intro + item bodies
    - engage2 build steps
    - engage button labels
    """

    raw_fragments: List[str] = []

    # Module-level text
    title = module.get("module_title")
    if title:
        raw_fragments.append(title)

    slides = module.get("slides", [])
    if not isinstance(slides, list):
        return normalize_and_filter(raw_fragments)

    for slide in slides:
        if not isinstance(slide, dict):
            continue

        # Header
        header = slide.get("header")
        if header:
            raw_fragments.append(header)

        # Notes
        notes = slide.get("notes")
        if notes:
            raw_fragments.append(notes)

        # Image (textual image identifiers count as text)
        image = slide.get("image")
        if image:
            raw_fragments.append(image)

        slide_type = slide.get("type")

        # -------------------------
        # PANEL
        # -------------------------
        if slide_type == "panel":
            content = slide.get("content", {})
            blocks = content.get("blocks", []) if isinstance(content, dict) else []

            for block in blocks:
                if not isinstance(block, dict):
                    continue
                btype = block.get("type")

                if btype == "paragraph":
                    text = block.get("text")
                    if text:
                        raw_fragments.append(text)

                elif btype == "bullets":
                    items = block.get("items", [])
                    for it in items:
                        if it:
                            raw_fragments.append(it)

        # -------------------------
        # ENGAGE 1
        # -------------------------
        elif slide_type == "engage":
            intro = slide.get("intro", {})
            if isinstance(intro, dict):
                raw_fragments.extend(_extract_text_from_text_block(intro))

            items = slide.get("items", [])
            for item in items:
                if not isinstance(item, dict):
                    continue

                label = item.get("button_label")
                if label:
                    raw_fragments.append(label)

                body = item.get("body", [])
                for block in body:
                    if not isinstance(block, dict):
                        continue
                    if block.get("type") == "paragraph":
                        txt = block.get("text")
                        if txt:
                            raw_fragments.append(txt)
                    elif block.get("type") == "bullets":
                        for it in block.get("items", []):
                            if it:
                                raw_fragments.append(it)

        # -------------------------
        # ENGAGE 2
        # -------------------------
        elif slide_type == "engage2":
            # Button labels
            labels = slide.get("button_labels", [])
            for lbl in labels:
                if lbl:
                    raw_fragments.append(lbl)

            build = slide.get("build", [])
            for step in build:
                if not isinstance(step, dict):
                    continue

                title = step.get("title")
                if title:
                    raw_fragments.append(title)

                content = step.get("content", [])
                for para in content:
                    if para:
                        raw_fragments.append(para)

    return normalize_and_filter(raw_fragments)


# -------------------------------------------------
# INTERNAL HELPERS
# -------------------------------------------------

def _extract_text_from_text_block(block: Dict[str, Any]) -> Iterable[str]:
    """
    Extracts text from a generic { title, content } block.
    """
    out: List[str] = []

    title = block.get("title")
    if title:
        out.append(title)

    content = block.get("content", [])
    for item in content:
        if item:
            out.append(item)

    return out

def extract_stage1_text_fragments(stage1_module: Dict[str, Any]) -> List[str]:
    """
    Extract ALL text fragments that Stage 1 claims to preserve.

    This mirrors Stage 1's responsibility exactly and matches
    the granularity used by Stage 1.1.
    """

    fragments: List[str] = []

    slides = stage1_module.get("slides", [])
    if not isinstance(slides, list):
        return fragments

    def extend_split(text: Any) -> None:
        if not text:
            return
        if not isinstance(text, str):
            text = str(text)
        for part in text.splitlines():
            s = part.replace("\u00A0", " ")
            s = " ".join(s.split()).strip()
            if s:
                fragments.append(s)

    for slide in slides:
        if not isinstance(slide, dict):
            continue

        extend_split(slide.get("header"))
        extend_split(slide.get("notes"))
        extend_split(slide.get("image"))

        content = slide.get("content", {})

        if slide.get("slide_type") == "panel":
            for b in content.get("blocks", []):
                if b.get("type") == "paragraph":
                    extend_split(b.get("text"))
                elif b.get("type") == "bullets":
                    for item in b.get("items", []):
                        extend_split(item)

        if slide.get("slide_type") == "engage1":
            intro = content.get("intro", {})
            if isinstance(intro, dict):
                extend_split(intro.get("text"))
                extend_split(intro.get("image"))
                extend_split(intro.get("notes"))

            for item in content.get("items", []):
                extend_split(item.get("button_label"))
                extend_split(item.get("image"))
                extend_split(item.get("notes"))

                for b in item.get("body", []):
                    if b.get("type") == "paragraph":
                        extend_split(b.get("text"))
                    elif b.get("type") == "bullets":
                        for it in b.get("items", []):
                            extend_split(it)

        if slide.get("slide_type") == "engage2":
            for lbl in content.get("button_labels", []):
                extend_split(lbl)

            for b in content.get("blocks", []):
                if b.get("type") == "paragraph":
                    extend_split(b.get("text"))
                elif b.get("type") == "bullets":
                    for it in b.get("items", []):
                        extend_split(it)

    return fragments
