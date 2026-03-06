from __future__ import annotations

import zipfile
import hashlib
from pathlib import Path
from typing import List, Set

from lxml import etree
from docx import Document

def is_column_header_text(text: str) -> bool:
    return text.lower() in {
        "image",
        "english text",
        "notes and instructions",
        "button labels",
    }

# WordprocessingML namespace
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
}


# -------------------------------------------------
# Normalization (must match audit normalization)
# -------------------------------------------------

def normalize_text(text: str) -> str:
    if not text:
        return ""

    # Normalize hyphen-like characters to standard ASCII hyphen
    text = (
        text.replace("\u2011", "-")  # non-breaking hyphen
            .replace("\u2013", "-")  # en dash
            .replace("\u2014", "-")  # em dash
            .replace("\u00AD", "")   # soft hyphen
            .replace("\u00A0", " ")  # non-breaking space
    )

    text = " ".join(text.split())
    return text.strip()


# -------------------------------------------------
# Slide header detection (MATCHES Stage 1)
# -------------------------------------------------

def is_slide_header_text(text: str) -> bool:
    t = text.lower()
    return (
        t.startswith("header:")
        or t.startswith("header ")
        or t.startswith("slide (header)")
        or t.startswith("slide header")
    )


# -------------------------------------------------
# XML helpers
# -------------------------------------------------

def paragraph_text(p: etree._Element) -> str:
    """
    Extract visible text from a <w:p> element.
    """
    texts = []
    for t in p.xpath(".//w:t", namespaces=NS):
        if t.text:
            texts.append(t.text)
    return normalize_text("".join(texts))


def xml_fingerprint(elem: etree._Element) -> str:
    """
    Create a stable hash of an XML element.
    This is how we dedupe Word's internal duplication.
    """
    raw = etree.tostring(elem, with_tail=False)
    return hashlib.sha256(raw).hexdigest()


# -------------------------------------------------
# Core extractor
# -------------------------------------------------

def extract_word_slide_text_fragments(docx_path: Path) -> List[str]:
    doc = Document(docx_path)

    fragments = []

    for table in doc.tables:
        is_slide_table = False

        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = normalize_text(p.text)
                    if txt and is_slide_header_text(txt):
                        is_slide_table = True
                        break

        if not is_slide_table:
            continue

        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = normalize_text(p.text)
                    if txt and not is_column_header_text(txt):
                        fragments.append(txt)

    return fragments
