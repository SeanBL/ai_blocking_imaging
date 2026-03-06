from __future__ import annotations

import json
from pathlib import Path
from typing import Dict, Any, List

from docx import Document

from .docx_formatting import (
    init_document_styles,
    add_slide_heading,
    add_review_table,
)
from .mapping import slide_to_rows


DEFAULT_INPUT = Path("data/review/module_review_suggestions.json")
DEFAULT_OUTPUT = Path("data/exports/module_editorial_review.docx")


def load_json(path: Path) -> Dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def export_review_docx(input_path: Path, output_path: Path) -> None:
    review = load_json(input_path)

    doc = Document()
    init_document_styles(doc)

    doc.add_heading("Editorial Review — LLM Suggestions", level=1)

    for slide in review.get("slides", []):
        slide_id = slide.get("slide_id")
        slide_type = slide.get("slide_type")

        # ✅ EXACT signature match
        add_slide_heading(doc, slide_id, slide_type)

        # mapping MUST return 6-tuples
        rows = slide_to_rows(slide)
        if rows:
            add_review_table(doc, rows)

        doc.add_paragraph("")  # spacing between slides

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    export_review_docx(DEFAULT_INPUT, DEFAULT_OUTPUT)


