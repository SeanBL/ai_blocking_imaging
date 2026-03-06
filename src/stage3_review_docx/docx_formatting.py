from __future__ import annotations
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor


BLUE = RGBColor(0x00, 0x70, 0xC0)
GRAY = RGBColor(0x66, 0x66, 0x66)


def init_document_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_slide_heading(doc: Document, slide_id: str, slide_type: str) -> None:
    h = doc.add_heading(f"{slide_id}  ({slide_type})", level=2)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def add_review_table(
    doc: Document,
    rows: list[tuple[str, str, str, str, str, str]]
) -> None:
    """
    rows:
      (block_id, block_type, original, suggested, notes, analysis)
    """

    table = doc.add_table(rows=0, cols=4)
    table.style = "Table Grid"

    for block_id, btype, original, suggested, notes, analysis in rows:
        # -------------------------------------------------
        # Metadata row (spans all columns)
        # -------------------------------------------------
        meta_row = table.add_row()
        meta_cell = meta_row.cells[0]

        # Merge across 4 columns
        for i in range(1, 4):
            meta_cell = meta_cell.merge(meta_row.cells[i])

        p = meta_cell.paragraphs[0]
        run = p.add_run(f"Block ID: {block_id}   |   Type: {btype}")
        run.bold = True
        run.font.color.rgb = GRAY

        # -------------------------------------------------
        # Header row
        # -------------------------------------------------
        hdr = table.add_row().cells
        hdr[0].text = "Original Text"
        hdr[1].text = "Suggested Revision"
        hdr[2].text = "Notes"
        hdr[3].text = "Analysis (Reviewer)"

        # -------------------------------------------------
        # Content row
        # -------------------------------------------------
        row = table.add_row().cells

        # Original
        row[0].paragraphs[0].add_run(original or "—")

        # Suggested
        p = row[1].paragraphs[0]
        run = p.add_run(suggested or "No change suggested.")
        run.font.color.rgb = BLUE

        # Notes
        p = row[2].paragraphs[0]
        run = p.add_run(notes or "")
        run.font.color.rgb = GRAY
        run.italic = True

        # Analysis
        p = row[3].paragraphs[0]
        run = p.add_run(analysis or "")
        run.font.color.rgb = GRAY
        run.italic = True
