from __future__ import annotations
from dataclasses import dataclass, asdict
from typing import List, Optional
from docx import Document
import json
import pathlib
from src.stage1.stage1_parse import parse_block
from src.stage1.models import RawBlock
import re


# ---------- DOCX loading ----------

def load_docx_in_order(path: str) -> List[str]:
    """
    Load a .docx file and return a flat list of text lines.

    This version:
    - Takes all top-level paragraphs
    - Then all paragraphs inside all tables
    It is simpler than walking the raw XML, but works well for
    WiRED-style modules where the main content is in tables.
    """
    doc = Document(path)
    lines: List[str] = []

    # 1) plain paragraphs (if any)
    for p in doc.paragraphs:
        text = p.text.strip()
        if text != "":
            lines.append(text)

    # 2) table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text.strip()
                    if text != "":
                        lines.append(text)

    return lines


# ---------- Header block segmentation ----------

def segment_header_blocks(lines: List[str]) -> List[RawBlock]:
    """
    Group lines into blocks, where each block starts at a line that begins with 'Header:'.

    Everything from 'Header:' up to (but not including) the next 'Header:' belongs
    to the same RawBlock.
    """
    blocks: List[RawBlock] = []
    current_header: Optional[str] = None
    current_lines: List[str] = []

    def flush_block():
        nonlocal current_header, current_lines, blocks
        if current_header is not None:
            # strip trailing empty lines
            while current_lines and current_lines[-1].strip() == "":
                current_lines.pop()
            blocks.append(RawBlock(header_line=current_header, lines=current_lines))
        current_header = None
        current_lines = []

    for raw_line in lines:
        line = raw_line or ""  # ensure not None

        if line.strip().startswith("Header:"):
            # start a new block
            flush_block()
            current_header = line.strip()  # e.g., "Header: Introduction"
            current_lines = []
        else:
            # accumulate lines under the current header if we are inside one
            if current_header is not None:
                current_lines.append(line)

    # flush last block
    flush_block()

    return blocks

# ---------- Remove duplicate blocks ----------
def filter_duplicate_blocks(blocks: List[RawBlock]) -> List[RawBlock]:
    """
    Only drop blocks that have no meaningful content at all.
    Keep all blocks that contain real text, even if headers repeat.
    """
    cleaned = []
    for block in blocks:
        has_meaningful = any(
            ln.strip() not in ("Image", "English text", "Translation text", "")
            for ln in block.lines
        )
        if has_meaningful:
            cleaned.append(block)
    return cleaned

# ---------- Main script ----------

def main():
    # Adjust this path to your actual docx file location
    root = pathlib.Path(__file__).resolve().parents[2]  # go up from src/
    docx_path = root / "data" / "raw" / "raw_module.docx"

    print(f"Loading: {docx_path}")
    lines = load_docx_in_order(str(docx_path))

    print(f"Total extracted lines: {len(lines)}")

    # Debug: print first 20 lines so we can see if 'Header:' shows up
    print("\nFirst 20 lines:")
    for l in lines[:20]:
        print("  ", repr(l))

    # -------- Stage 1 Part 1: Extract Raw Blocks --------
    blocks = segment_header_blocks(lines)
    blocks = filter_duplicate_blocks(blocks)
    print(f"\nDetected header blocks (after dedupe): {len(blocks)}")

    # -------- Stage 1 Part 2: Parse Blocks --------

    parsed_blocks = [parse_block(b) for b in blocks]

    parsed_dir = root / "data" / "processed" / "stage1_blocks"
    parsed_dir.mkdir(parents=True, exist_ok=True)

    def safe_filename(text: str) -> str:
        t = text.strip().replace(" ", "_")
        t = re.sub(r"[^\w\-]", "", t)
        return t

    for idx, pb in enumerate(parsed_blocks, start=1):
        # create safe filename
        safe_header = (
            pb.header
            .replace(" ", "_")
            .replace("/", "_")
            .replace(":", "")
        )
        filename = f"{idx:03d}_{safe_header}.json"
        out_path = parsed_dir / filename

        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(asdict(pb), f, ensure_ascii=False, indent=2)

    print(f"\nSaved {len(parsed_blocks)} parsed blocks to folder:")
    print(parsed_dir)

    # -------- PREVIEW --------
    print("\nPreview of first 3 parsed blocks:")
    for i, pb in enumerate(parsed_blocks[:3], start=1):
        print("=" * 60)
        print(f"Parsed Block {i}: {pb.header}")
        print("Image Raw:", pb.image_raw)
        print("Block Type Raw:", pb.block_type_raw)
        print("Notes Raw:", pb.notes_raw)
        print("English (first 200 chars):")
        if pb.english_text_raw:
            print("English (first paragraph):", pb.english_text_raw[0][:200], "...")
        else:
            print("No English text found.")

        print()

    # -------- SAVE RAW BLOCKS (Part 1 OUTPUT) --------
    out_raw = root / "data" / "processed" / "stage1_raw_blocks.json"
    out_raw.parent.mkdir(parents=True, exist_ok=True)
    with open(out_raw, "w", encoding="utf-8") as f:
        json.dump([asdict(b) for b in blocks], f, ensure_ascii=False, indent=2)
    print(f"Raw blocks written to: {out_raw}")

    # -------- SAVE PARSED BLOCKS (Part 2 OUTPUT) --------
    out_parsed = root / "data" / "processed" / "stage1_parsed_blocks.json"
    with open(out_parsed, "w", encoding="utf-8") as f:
        json.dump([asdict(pb) for pb in parsed_blocks], f, ensure_ascii=False, indent=2)
    print(f"Parsed blocks written to: {out_parsed}")


if __name__ == "__main__":
    main()