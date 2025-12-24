from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document


# -----------------------------
# Regex helpers
# -----------------------------
RE_SLIDE_TYPE = re.compile(r"Slide\s*Type\s*=\s*(Panel|Engage\s*1|Engage\s*2)", re.IGNORECASE)

# =============================
# Explicit quiz markers (v2)
# =============================

# [[QUIZ:3]]
RE_QUIZ_START = re.compile(r"\[\[\s*QUIZ\s*:\s*(\d+)\s*\]\]", re.IGNORECASE)

# [[QUIZ:3:QUESTIONS=6]]
RE_QUIZ_END = re.compile(
    r"\[\[\s*QUIZ\s*:\s*(\d+)\s*:\s*QUESTIONS\s*=\s*(\d+)\s*\]\]",
    re.IGNORECASE
)

def normalize_ws(s: str) -> str:
    # Preserve newlines, normalize internal whitespace per line
    lines = [re.sub(r"[ \t]+", " ", ln).strip() for ln in s.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    # Keep blank lines (they can be meaningful separators)
    return "\n".join(lines).strip()


def cell_text(cell) -> str:
    # docx cell.text flattens; we still use it for header/image/notes/labels.
    return normalize_ws(cell.text or "")


def is_bullet_paragraph(p) -> bool:
    """
    Robust-ish bullet detection:
    - True if paragraph has numbering properties (numPr), which docx uses for lists.
    This catches both bulleted and numbered lists.
    """
    try:
        pPr = p._p.pPr  # type: ignore[attr-defined]
        if pPr is None:
            return False
        return pPr.numPr is not None  # bullet/numbered list
    except Exception:
        return False
    
def text_to_blocks(text: str) -> List[Dict[str, Any]]:
    """
    Convert linearized English Text into paragraph blocks.
    Bullet fidelity is intentionally deferred to a later stage.
    """
    if not text:
        return []

    lines = [
        ln.strip()
        for ln in text.split("\n")
        if ln.strip()
    ]

    return [
        {"type": "paragraph", "text": ln}
        for ln in lines
    ]

def extract_blocks_from_cell(cell) -> List[Dict[str, Any]]:
    """
    Extract content.blocks preserving paragraph vs bullets.

    Output example:
    [
      {"type": "paragraph", "text": "..."},
      {"type": "bullets", "items": ["...", "..."]},
      {"type": "paragraph", "text": "..."}
    ]
    """
    blocks: List[Dict[str, Any]] = []
    bullet_acc: List[str] = []

    def flush_bullets():
        nonlocal bullet_acc
        if bullet_acc:
            blocks.append({"type": "bullets", "items": bullet_acc})
            bullet_acc = []

    for p in cell.paragraphs:
        raw = (p.text or "").strip()
        # Keep true blank paragraphs as separators (for engage segmentation later),
        # but do not emit them as blocks.
        if not raw:
            flush_bullets()
            continue

        if is_bullet_paragraph(p):
            bullet_acc.append(normalize_ws(raw))
        else:
            flush_bullets()
            blocks.append({"type": "paragraph", "text": normalize_ws(raw)})

    flush_bullets()
    return blocks


def parse_slide_type(notes: str) -> str:
    m = RE_SLIDE_TYPE.search(notes or "")
    if not m:
        raise ValueError("Missing required 'Slide Type = ...' in Notes/Instructions.")
    val = m.group(1).lower().replace(" ", "")
    if val == "panel":
        return "panel"
    if val == "engage1":
        return "engage1"
    if val == "engage2":
        return "engage2"
    raise ValueError(f"Unrecognized slide type: {m.group(1)}")

def extract_slide_type(*fields: str) -> Optional[str]:
    for text in fields:
        if not text:
            continue
        m = RE_SLIDE_TYPE.search(text)
        if m:
            val = m.group(1).lower().replace(" ", "")
            if val in ("panel", "engage1", "engage2"):
                return val
    return None 

def split_button_labels(raw: str) -> List[str]:
    if not raw:
        return []
    lines = [ln.strip() for ln in raw.split("\n")]
    return [ln for ln in lines if ln]


def blocks_to_plaintext(blocks: List[Dict[str, Any]]) -> str:
    """
    Utility for simple engage prebuilt segmentation:
    - Paragraphs become lines
    - Bullets become "- item" lines
    """
    out: List[str] = []
    for b in blocks:
        if b["type"] == "paragraph":
            out.append(b["text"])
        elif b["type"] == "bullets":
            out.extend([f"- {it}" for it in b["items"]])
    return "\n".join(out).strip()


def segment_paragraph_blocks(blocks: List[Dict[str, Any]]) -> List[str]:
    """
    Very simple segmentation for prebuilt engages:
    - Use paragraph blocks only (ignore bullets) as segments, in order.
    - This matches the common pattern: intro paragraph + item paragraphs.
    """
    return [b["text"] for b in blocks if b.get("type") == "paragraph" and b.get("text")]


@dataclass
class QuizState:
    quiz_number: int
    question_count: Optional[int] = None
    source_slide_ids: List[str] = None

    def __post_init__(self):
        if self.source_slide_ids is None:
            self.source_slide_ids = []


def extract_tables_as_records(doc: Document) -> List[Dict[str, Any]]:
    LABELS = {
        "slide (header)": "Slide (Header)",
        "english text": "English Text",
        "notes and instructions": "Notes and Instructions",
        "button labels": "Button Labels",
        "image": "Image",
    }

    def normalize(s: str) -> str:
        s = s.replace("\u00A0", " ")
        s = re.sub(r"\s+", " ", s)
        return s.strip()

    def norm(s: str) -> str:
        return normalize(s).lower()

    records: List[Dict[str, str]] = []

    for tbl in doc.tables:
        record: Dict[str, str] = {"_slide_type_raw": ""}
        current_label: Optional[str] = None
        buffer: List[str] = []

        def flush():
            nonlocal buffer, current_label
            if current_label and buffer:
                record[current_label] = normalize(" ".join(buffer))
            buffer = []

        # 🔑 SINGLE-CELL, ROW-BASED PARSING
        for row in tbl.rows:
            # concatenate all cells in the row (some rows have merged cells)
            row_text = normalize(" ".join(cell.text for cell in row.cells))
            if not row_text:
                continue

            # capture slide type anywhere
            m = RE_SLIDE_TYPE.search(row_text)
            if m:
                record["_slide_type_raw"] = m.group(0)

            row_norm = norm(row_text)

            # label switch (EXACT match only)
            label_match = None
            for key_norm, canonical in LABELS.items():
                if row_norm == key_norm:
                    label_match = canonical
                    break

            if label_match:
                flush()
                current_label = label_match
                continue

            # content line
            if current_label:
                buffer.append(row_text)

        flush()

        # only keep real slides
        if any(k for k in record if not k.startswith("_")):
            records.append(record)

    return records




def stage1_extract_v2(docx_path: Path) -> Dict[str, Any]:
    doc = Document(str(docx_path))
    tables = extract_tables_as_records(doc)

    module: Dict[str, Any] = {
        "module_title": docx_path.stem,
        "slides": []
    }

    last_seen_header: Optional[str] = None
    last_seen_slide_type: Optional[str] = None
    slide_index = 0

    active_quiz: Optional[QuizState] = None

    for rec in tables:
        slide_index += 1
        slide_id = f"slide_{slide_index:03d}"

        # Extract raw fields
        header_raw = normalize_ws(rec.get("Slide (Header)", ""))
        image_raw = normalize_ws(rec.get("Image", ""))
        notes_raw = normalize_ws(
            rec.get("Notes and Instructions")
            or rec.get("Notes & Instructions")
            or rec.get("Notes", "")
        )
        labels_raw = normalize_ws(rec.get("Button Labels", ""))

        raw_slide_type = extract_slide_type(
            rec.get("_slide_type_raw", ""),
            notes_raw,
            rec.get("English Text", ""),
            header_raw
        )

        if raw_slide_type:
            slide_type = raw_slide_type
            last_seen_slide_type = slide_type
        else:
            if last_seen_slide_type is None:
                # FIRST SLIDE DEFAULT
                slide_type = "panel"
                last_seen_slide_type = slide_type
            else:
                slide_type = last_seen_slide_type

        # Header inheritance rule (only panels can omit header)
        if header_raw:
            header = header_raw
            last_seen_header = header_raw
        else:
            if slide_type == "panel":
                if last_seen_header is None:
                    # FIRST SLIDE DEFAULT HEADER
                    header = module["module_title"]
                    last_seen_header = header
                else:
                    header = f"{last_seen_header} (Continued)"
            else:
                raise ValueError(
                    f"Non-panel slide missing header (slide_type={slide_type}) at table #{slide_index}."
                )

        # Blocks (preserve bullets)
        english_raw = rec.get("English Text", "")
        blocks = text_to_blocks(english_raw)

        # Button labels / prebuilt flag
        button_labels = split_button_labels(labels_raw)
        prebuilt = len(button_labels) > 0

        # Build slide object (base)
        slide_obj: Dict[str, Any] = {
            "id": slide_id,
            "header": header,
            "slide_type": slide_type,
            "image": image_raw if image_raw else None,
            "notes": notes_raw,
            "content": {
                "blocks": blocks
            }
        }

        # Attach engage metadata if relevant
        if slide_type in ("engage1", "engage2"):
            slide_obj["content"]["prebuilt"] = prebuilt
            if button_labels:
                slide_obj["content"]["button_labels"] = button_labels

            # If engage1 is prebuilt, attempt a light structural extraction:
            # intro = first paragraph; items = subsequent paragraphs mapped to labels
            if slide_type == "engage1" and prebuilt:
                paras = segment_paragraph_blocks(blocks)
                if paras:
                    intro_text = paras[0]
                    item_texts = paras[1:]
                    # Map labels to items if counts line up; otherwise keep raw blocks + labels
                    if len(item_texts) == len(button_labels) and len(button_labels) > 0:
                        slide_obj["content"]["intro"] = {"text": intro_text}
                        slide_obj["content"]["items"] = [
                            {"label": button_labels[i], "text": item_texts[i]}
                            for i in range(len(button_labels))
                        ]
                    else:
                        # Still mark prebuilt, but include a warning for human review.
                        slide_obj["content"]["prebuilt_warning"] = (
                            f"Could not map {len(button_labels)} button labels to {len(item_texts)} item paragraphs. "
                            f"Preserved raw blocks and button_labels."
                        )

            # Engage2: store a single button label if present (optional)
            if slide_type == "engage2" and button_labels:
                # If the author put multiple lines, keep them as-is, but typically you want one.
                slide_obj["content"]["button_label"] = button_labels[0]

        # -----------------------------
        # Quiz range handling (inclusive)
        # -----------------------------
        quiz_text = "\n".join([
            notes_raw,
            rec.get("English Text", ""),
            header_raw,
        ])
        # Detect end marker first (in case both appear in same notes).
        end_m = RE_QUIZ_END.search(quiz_text)
        start_m = RE_QUIZ_START.search(quiz_text)

        # If we're currently inside an active quiz range, include this slide.
        if active_quiz is not None:
            active_quiz.source_slide_ids.append(slide_id)

        # If a start marker appears and no active quiz yet, open range and include this slide.
        if start_m is not None:
            quiz_num = int(start_m.group(1))
            if active_quiz is not None:
                raise ValueError(
                    f"Encountered QUIZ {quiz_num} start while Quiz {active_quiz.quiz_number} is still active."
                )
            active_quiz = QuizState(quiz_number=quiz_num)
            active_quiz.source_slide_ids.append(slide_id)

        # Add slide to module (slides always preserved)
        module["slides"].append(slide_obj)

        # If end marker appears, close range (inclusive) and emit quiz placeholder.
        if end_m is not None:
            quiz_num_end = int(end_m.group(1))
            q_count = int(end_m.group(2))

            # If no active quiz, open-and-close on same slide (still inclusive)
            if active_quiz is None:
                active_quiz = QuizState(quiz_number=quiz_num_end)
                active_quiz.source_slide_ids.append(slide_id)

            if active_quiz.quiz_number != quiz_num_end:
                raise ValueError(
                    f"Quiz end marker 'Quiz {quiz_num_end} Questions = {q_count}' "
                    f"does not match active Quiz {active_quiz.quiz_number}."
                )

            active_quiz.question_count = q_count

            quiz_id = f"quiz_{quiz_num_end:02d}"
            quiz_obj: Dict[str, Any] = {
                "id": quiz_id,
                "slide_type": "quiz",
                "quiz_number": quiz_num_end,
                "question_count": q_count,
                "source_slide_ids": active_quiz.source_slide_ids,
                "questions": []
            }

            module["slides"].append(quiz_obj)
            active_quiz = None

    if active_quiz is not None:
        raise ValueError(
            f"Quiz {active_quiz.quiz_number} started but never ended with 'Quiz {active_quiz.quiz_number} Questions = N'."
        )

    return module


def main() -> None:
    parser = argparse.ArgumentParser(description="Stage 1 v2 extractor: Word .docx -> module_v2.json")
    parser.add_argument("--in", dest="in_path", required=True, help="Path to input .docx file")
    parser.add_argument("--out", dest="out_path", required=True, help="Path to output .json file")
    args = parser.parse_args()

    in_path = Path(args.in_path).expanduser().resolve()
    out_path = Path(args.out_path).expanduser().resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)

    module = stage1_extract_v2(in_path)
    out_path.write_text(json.dumps(module, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"✅ Wrote: {out_path}")


if __name__ == "__main__":
    main()
