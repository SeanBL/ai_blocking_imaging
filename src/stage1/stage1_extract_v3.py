"""
STAGE 1 INVARIANT CONTRACT (DO NOT VIOLATE)

Stage 1 is a STRUCTURAL extractor only.

INVARIANTS:
1. Paragraphs in Word are preserved as paragraphs in JSON.
2. Bulleted / numbered lists explicitly marked by Word are preserved as bullets.
3. Mixed prose + bullets in the same cell are preserved in order.
4. Stage 1 NEVER infers bullets, engages, or pedagogy.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Dict, List, Any, Optional

from docx import Document
from docx.oxml.ns import qn


# -------------------------------
# Utilities
# -------------------------------

def normalize(text: str) -> str:
    if not text:
        return ""
    return " ".join(text.replace("\u00A0", " ").split()).strip()


def is_list_paragraph(p) -> bool:
    """
    True if the paragraph is explicitly a Word list item (bullet or numbered).
    This is STRUCTURAL detection, not inference.
    """
    pPr = p._p.pPr
    if pPr is None:
        return False
    return pPr.find(qn("w:numPr")) is not None


def canonical_col_label(raw: str) -> str:
    t = normalize(raw).lower()
    if "/" in t:
        t = t.split("/", 1)[0].strip()

    if t.startswith("notes and instructions"):
        return "notes and instructions"
    if t.startswith("english text"):
        return "english text"
    if t.startswith("image"):
        return "image"
    if t.startswith("button labels"):
        return "button labels"

    return t


def is_slide_header(text: str) -> bool:
    t = text.lower()
    return (
        t.startswith("header:")
        or t.startswith("header ")
        or t.startswith("slide (header)")
        or t.startswith("slide header")
    )


# -------------------------------
# Extractor
# -------------------------------

def extract_tables_v3(docx_path: Path) -> Dict[str, Any]:
    doc = Document(str(docx_path))

    module: Dict[str, Any] = {"module_title": docx_path.stem, "slides": []}
    slide_index = 0

    slide: Optional[Dict[str, Any]] = None
    col_labels: List[str] = []
    columns: Dict[str, List[str]] = {}

    # Engage1 state
    engage1_mode = False
    engage1_intro_row_pending = False
    collecting_button_labels = False

    # Engage2 state
    engage2_mode = False

    engage_intro_parts: List[str] = []
    engage_intro_image: Optional[str] = None
    engage_intro_notes: List[str] = []
    engage_items: List[Dict[str, Any]] = []

    def reset_engage_state():
        nonlocal engage1_mode, engage1_intro_row_pending, collecting_button_labels
        nonlocal engage2_mode
        nonlocal engage_intro_parts, engage_intro_image, engage_intro_notes, engage_items

        engage1_mode = False
        engage2_mode = False
        engage1_intro_row_pending = False
        collecting_button_labels = False
        engage_intro_parts = []
        engage_intro_image = None
        engage_intro_notes = []
        engage_items = []

    def finalize_slide(cur_slide: Dict[str, Any], cur_columns: Dict[str, List[str]]) -> None:
        notes = "\n".join(cur_columns.get("notes and instructions", [])).strip()
        cur_slide["notes"] = notes

        notes_lower = notes.lower()
        if "slide type = engage 1" in notes_lower:
            cur_slide["slide_type"] = "engage1"
        elif "slide type = engage 2" in notes_lower:
            cur_slide["slide_type"] = "engage2"
        else:
            cur_slide["slide_type"] = "panel"

        # ------------------------------------------------
        # IMAGE COLUMN ROUTING (STRUCTURAL — Stage 1)
        #
        # For non-Engage1 slides, the Image column represents the slide-level image field.
        # (Engage1 uses per-intro/per-item images and should NOT be overwritten here.)
        # ------------------------------------------------
        if cur_slide["slide_type"] != "engage1":
            img_lines = []
            for t in cur_columns.get("image", []):
                tt = normalize(t)
                # Avoid treating the Engage2 "Button Labels" sentinel as an image value.
                if tt and tt.lower() != "button labels":
                    img_lines.append(tt)
            cur_slide["image"] = "\n".join(img_lines).strip() or None

        if cur_slide["slide_type"] == "engage1":
            labels: List[str] = []
            for txt in cur_columns.get("button labels", []):
                labels.extend([p.strip() for p in txt.split("|") if p.strip()])

            for i, item in enumerate(engage_items):
                if i < len(labels):
                    item["button_label"] = labels[i]

            cur_slide["image"] = None
            cur_slide["content"] = {
                "intro": {
                    "text": "\n".join(engage_intro_parts).strip(),
                    "image": engage_intro_image,
                    "notes": "\n".join(engage_intro_notes).strip() or None,
                },
                "items": engage_items,
            }
            return

        # -------------------------------
        # ENGAGE 2 FINALIZATION
        # -------------------------------
        if cur_slide["slide_type"] == "engage2":
            labels: List[str] = []
            for txt in cur_columns.get("button labels", []):
                labels.extend([p.strip() for p in txt.split("|") if p.strip()])

            if labels:
                cur_slide.setdefault("content", {})
                cur_slide["content"]["button_labels"] = labels

            # ✅ SCHEMA GUARD — ADD THIS EXACTLY HERE
            cur_slide.setdefault("content", {})
            cur_slide["content"].setdefault("button_labels", [])

    def start_new_slide(header_text: str, tbl, header_row_idx: int) -> None:
        nonlocal slide, columns, col_labels, slide_index

        if slide is not None:
            finalize_slide(slide, columns)
            module["slides"].append(slide)

        slide_index += 1
        slide = {
            "id": f"slide_{slide_index:03d}",
            "header": header_text,
            "slide_type": "panel",
            "image": None,
            "notes": "",
            "content": {"blocks": []},
        }

        reset_engage_state()

        label_row = tbl.rows[header_row_idx + 1]
        col_labels = [canonical_col_label(c.text) for c in label_row.cells]
        columns = {label: [] for label in col_labels if label}

    # -------------------------------
    # Main table walk
    # -------------------------------

    for tbl in doc.tables:
        if len(tbl.rows) < 3:
            continue

        row_idx = 0

        while row_idx < len(tbl.rows):
            row = tbl.rows[row_idx]
            first_cell = normalize(row.cells[0].text)

            if is_slide_header(first_cell):
                start_new_slide(first_cell, tbl, row_idx)
                row_idx += 2
                continue

            if slide is None:
                row_idx += 1
                continue

            # -------------------------------
            # Row-level extraction
            # -------------------------------
            row_texts: Dict[str, List[str]] = {}

            for idx, cell in enumerate(row.cells):
                if idx >= len(col_labels):
                    continue
                label = col_labels[idx]
                if not label:
                    continue

                texts = [normalize(p.text) for p in cell.paragraphs if normalize(p.text)]
                if texts:
                    row_texts[label] = texts
                    columns.setdefault(label, []).extend(texts)

            img_texts = row_texts.get("image", [])
            eng_texts = row_texts.get("english text", [])
            notes_texts = row_texts.get("notes and instructions", [])

            # Engage detection
            if notes_texts:
                blob = " ".join(t.lower() for t in notes_texts)
                if "slide type = engage 1" in blob:
                    engage1_mode = True
                    engage2_mode = False
                    engage1_intro_row_pending = True
                    collecting_button_labels = False
                elif "slide type = engage 2" in blob:
                    engage2_mode = True
                    engage1_mode = False

            # -------------------------------
            # ENGAGE 1 parsing
            # -------------------------------
            if engage1_mode:
                if img_texts and img_texts[0].strip().lower() == "button labels":
                    collecting_button_labels = True
                    if eng_texts:
                        columns.setdefault("button labels", []).extend(eng_texts)
                    row_idx += 1
                    continue

                if collecting_button_labels:
                    if eng_texts:
                        columns.setdefault("button labels", []).extend(eng_texts)
                    row_idx += 1
                    continue

                if engage1_intro_row_pending:
                    if img_texts:
                        engage_intro_image = " ".join(img_texts).strip()
                    if eng_texts:
                        engage_intro_parts.extend(eng_texts)
                    if notes_texts:
                        engage_intro_notes.extend(notes_texts)
                    engage1_intro_row_pending = False
                    row_idx += 1
                    continue

                if eng_texts:
                    eng_blocks: List[Dict[str, Any]] = []
                    eng_col_idx = col_labels.index("english text")
                    eng_cell = row.cells[eng_col_idx]

                    paras = [p for p in eng_cell.paragraphs if normalize(p.text)]
                    list_items: List[str] = []

                    for p in paras:
                        txt = normalize(p.text)
                        if is_list_paragraph(p):
                            list_items.append(txt)
                        else:
                            if list_items:
                                eng_blocks.append({"type": "bullets", "items": list_items})
                                list_items = []
                            eng_blocks.append({"type": "paragraph", "text": txt})

                    if list_items:
                        eng_blocks.append({"type": "bullets", "items": list_items})

                    engage_items.append({
                        "button_label": None,
                        "image": " ".join(img_texts).strip() if img_texts else None,
                        "body": eng_blocks,
                        "notes": "\n".join(notes_texts).strip() if notes_texts else None,
                    })

                    row_idx += 1
                    continue

                # FIX 3 — engage rows NEVER fall through
                row_idx += 1
                continue

            # -------------------------------
            # ENGAGE 2 button label row
            # -------------------------------
            if engage2_mode:
                if img_texts and img_texts[0].strip().lower() == "button labels":
                    # STRUCTURAL: button labels live in their own column
                    if eng_texts:
                        slide.setdefault("content", {})
                        slide["content"].setdefault("button_labels", [])
                        slide["content"]["button_labels"].extend(eng_texts)

                    row_idx += 1
                    continue

            # -------------------------------
            # PANEL parsing (single-pass)
            # -------------------------------
            blocks = slide["content"]["blocks"]

            if eng_texts:
                eng_col_idx = col_labels.index("english text")
                eng_cell = row.cells[eng_col_idx]

                for p in eng_cell.paragraphs:
                    txt = normalize(p.text)
                    if not txt:
                        continue

                    if is_list_paragraph(p):
                        if blocks and blocks[-1]["type"] == "bullets":
                            blocks[-1]["items"].append(txt)
                        else:
                            blocks.append({"type": "bullets", "items": [txt]})
                    else:
                        blocks.append({"type": "paragraph", "text": txt})

            row_idx += 1

    if slide is not None:
        finalize_slide(slide, columns)
        module["slides"].append(slide)

    return module


# -------------------------------
# CLI
# -------------------------------

def main() -> None:
    import argparse

    parser = argparse.ArgumentParser(description="Stage 1 v3 extractor (final)")
    parser.add_argument("--in", dest="in_path", required=False)
    parser.add_argument("--out", dest="out_path", required=False)
    args = parser.parse_args()

    raw_dir = Path("data/raw")

    if not args.in_path:
        docs = list(raw_dir.glob("*.docx"))
        if not docs:
            raise SystemExit("❌ No Word documents found in data/raw")
        if len(docs) > 1:
            raise SystemExit(f"❌ Multiple Word documents found: {[d.name for d in docs]}")
        input_path = docs[0]
    else:
        input_path = Path(args.in_path)

    output_path = Path(args.out_path) if args.out_path else Path("data/processed/module_v3.json")

    module = extract_tables_v3(input_path)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(module, indent=2, ensure_ascii=False), encoding="utf-8")

    print(f"✅ Input: {input_path.name}")
    print(f"✅ Output: {output_path}")
    print(f"✅ Slides: {len(module['slides'])}")


if __name__ == "__main__":
    main()
