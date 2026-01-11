from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Dict, List

from docx import Document

from .docx_formatting import (
    init_document_styles,
    add_slide_table,
)
from .mapping import slide_to_table_rows


DEFAULT_INPUT = Path("data/processed/module_stage2_9.json")
DEFAULT_OUTPUT = Path("data/exports/module_stage3B.docx")


def load_json(path: Path) -> Dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def export_module_to_docx(module_json: Dict[str, Any], out_path: Path) -> None:
    doc = Document()
    init_document_styles(doc)

    slides: List[Dict[str, Any]] = module_json.get("slides", [])
    module_title = module_json.get("module_title") or module_json.get("module_id") or "Module"

    # Optional title at top (can be removed if you want the tables only)
    doc.add_paragraph(module_title)

    # --- Reorder slides for authoring clarity ---
    instructional_slides: List[Dict[str, Any]] = []
    quiz_application_slides: List[Dict[str, Any]] = []
    quiz_final_slides: List[Dict[str, Any]] = []

    for slide in slides:
        slide_type = (slide.get("slide_type") or slide.get("type") or "").lower()
        slide_id = (slide.get("id") or "").lower()

        if slide_type == "quiz":
            if slide_id.endswith("_application"):
                quiz_application_slides.append(slide)
            elif slide_id.endswith("_final"):
                quiz_final_slides.append(slide)
            else:
                # fallback: keep inline quizzes with instructional flow
                instructional_slides.append(slide)
        else:
            instructional_slides.append(slide)

    ordered_slides = (
        instructional_slides
        + quiz_application_slides
        + quiz_final_slides
    )

    # --- Render slides in new order ---
    for slide_index, slide in enumerate(ordered_slides, start=1):
        header_text = (
            slide.get("header")
            or slide.get("title")
            or slide.get("id")
            or f"Slide {slide_index:03d}"
        )

        rows = slide_to_table_rows(slide)

        add_slide_table(
            doc=doc,
            header_text=header_text,
            rows=rows,
        )

        doc.add_paragraph("")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    module_path = DEFAULT_INPUT
    out_path = DEFAULT_OUTPUT

    module_json = load_json(module_path)
    export_module_to_docx(module_json, out_path)
    print(f"✅ Wrote: {out_path}")


if __name__ == "__main__":
    main()
