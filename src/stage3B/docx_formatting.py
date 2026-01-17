from __future__ import annotations

from typing import List, Tuple

from docx.document import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.shared import RGBColor
from docx.shared import Pt

BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)

# One row = (image_cell_text, english_text, notes_text)
RowTriple = Tuple[str, str, str]

# Quiz row = (english_question, english_answer, translated_question, translated_answer)
QuizRow = Tuple[str, str, str, str]


def init_document_styles(doc: Document) -> None:
    # Keep it simple + deterministic (you can expand later)
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def _set_cell_shading(cell, fill_hex: str) -> None:
    """
    Deterministic header shading without relying on Word themes.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)

def _is_multiline_text(text: str) -> bool:
    return "\n" in text

def _render_multiline_text(cell, text: str) -> None:
    """
    Render newline-delimited text as controlled paragraphs
    inside a single table cell.
    """
    cell.text = ""

    parts = [p.strip() for p in text.split("\n") if p.strip()]
    if not parts:
        return

    # First paragraph (no top spacing)
    p = cell.paragraphs[0]
    p.text = parts[0]

    # Subsequent paragraphs (explicit spacing)
    for part in parts[1:]:
        p = cell.add_paragraph(part)
        p.paragraph_format.space_before = Pt(6)

def _apply_between_paragraph_spacing(p) -> None:
    # Space before works reliably inside tables, but only for non-first paragraphs
    p.paragraph_format.space_before = Pt(8)

def _write_paragraphs_to_cell(cell, parts: List[str]) -> None:
    """
    Writes paragraphs into a table cell WITHOUT creating top padding.
    Uses the cell's existing first paragraph for the first text chunk,
    then adds additional paragraphs with space_before.
    """
    cell.text = ""  # leaves one empty paragraph behind (Word behavior)

    if not parts:
        return

    # Reuse the existing first paragraph so there's no spacing above it
    p0 = cell.paragraphs[0]
    p0.text = parts[0]

    # Remaining paragraphs get spacing before (between paragraphs)
    for part in parts[1:]:
        p = cell.add_paragraph(part)
        _apply_between_paragraph_spacing(p)

def add_slide_table(doc: Document, header_text: str, rows: List[RowTriple]) -> None:
    """
    Creates one table per slide:
      Row 0: merged header
      Row 1: column labels
      Row 2+: content rows
    """
    # +2 for header + labels
    table = doc.add_table(rows=len(rows) + 2, cols=3)
    table.style = "Table Grid"

    # Column widths (tweak to match your authoring doc)
    col_widths = [Inches(1.4), Inches(4.6), Inches(2.0)]
    for col_idx, w in enumerate(col_widths):
        for r in table.rows:
            r.cells[col_idx].width = w

    # Row 0: merged header
    hdr = table.rows[0].cells
    merged = hdr[0].merge(hdr[1]).merge(hdr[2])
    p = merged.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(header_text)
    run.bold = True
    run.font.color.rgb = BLUE
    _set_cell_shading(merged, "D9D9D9")

    # Row 1: labels
    labels = table.rows[1].cells
    labels[0].text = "Image"
    labels[1].text = "English Text"
    labels[2].text = "Notes and Instructions"
    for c in labels:
        _set_cell_shading(c, "EFEFEF")
        for run in c.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = BLUE

    # Content rows
    for i, (img_txt, eng_txt, notes_txt) in enumerate(rows, start=2):
        cells = table.rows[i].cells
        cells[0].text = img_txt or ""
        cell = cells[1]
        cell.text = ""  # clear default paragraph

        if isinstance(eng_txt, list):
            # Use the cell's first paragraph for the first line to avoid top padding
            cell.text = ""
            first = True

            for text, is_bullet in eng_txt:
                if first:
                    p = cell.paragraphs[0]
                    p.text = text
                    first = False
                else:
                    p = cell.add_paragraph(text)
                    _apply_between_paragraph_spacing(p)

                if is_bullet:
                    p.style = "List Bullet"

        else:
            if isinstance(eng_txt, str) and "\n" in eng_txt:
                _render_multiline_text(cell, eng_txt)
            elif eng_txt:
                _write_paragraphs_to_cell(cell, [str(eng_txt)])

        notes_cell = cells[2]

        if notes_txt:
            # Use first paragraph to avoid top padding
            notes_cell.text = ""
            p = notes_cell.paragraphs[0]

            if "Slide Type = Engage 1" in notes_txt:
                before, after = notes_txt.split("Slide Type = Engage 1", 1)

                if before.strip():
                    p.add_run(before)

                r = p.add_run("Slide Type = Engage 1")
                r.font.color.rgb = RED
                r.bold = True

                if after.strip():
                    p.add_run(after)
            else:
                p.add_run(notes_txt)

def add_quiz_table(doc: Document, quiz_rows: List[Tuple[str, str, str, str]]) -> None:
    table = doc.add_table(rows=len(quiz_rows) + 1, cols=4)
    table.style = "Table Grid"

    headers = [
        "English question",
        "English answer",
        "Translated question",
        "Translated answer",
    ]

    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = h
        _set_cell_shading(cell, "EFEFEF")
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = BLUE

    for row_idx, row in enumerate(quiz_rows, start=1):
        for col_idx, text in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            if text:
                parts = [p.strip() for p in str(text).split("\n\n") if p.strip()]
                if parts:
                    p0 = cell.paragraphs[0]
                    p0.text = parts[0]
                    for part in parts[1:]:
                        p = cell.add_paragraph(part)
                        _apply_between_paragraph_spacing(p)


